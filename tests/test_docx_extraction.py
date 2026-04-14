import pytest
from pathlib import Path
import tempfile

# If python-docx is not available, skip this test early
pytest.importorskip("docx")
from docx import Document

from db_optimized import _docx_to_text_one


def test_docx_to_text_one_creates_text_and_meta():
    with tempfile.TemporaryDirectory() as td:
        p = Path(td) / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello world.")
        doc.add_paragraph("Second paragraph.")
        doc.save(str(p))

        text, meta = _docx_to_text_one(p)

        assert isinstance(text, str)
        assert "Hello world." in text
        assert "Second paragraph." in text
        assert isinstance(meta, dict)
        assert meta.get("source") == p.name

